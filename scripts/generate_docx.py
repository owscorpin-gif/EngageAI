#!/usr/bin/env python3
"""
Engage AI — Professional Project Description Generator
Produces a richly formatted .docx document with embedded screenshots.
"""

import os
import sys
from pathlib import Path
from datetime import date

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌  python-docx not found. Run: pip3 install python-docx pillow")
    sys.exit(1)

# ─── Paths ───────────────────────────────────────────────────────────────────
BASE   = Path(__file__).parent.parent
IMAGES = BASE / "docs_images"
OUTPUT = BASE / "Engage_AI_Project_Description.docx"

# ─── Brand colours ───────────────────────────────────────────────────────────
INDIGO   = RGBColor(0x4F, 0x46, 0xE5)   # primary-600
VIOLET   = RGBColor(0x7C, 0x3A, 0xED)   # accent-600
SLATE900 = RGBColor(0x0F, 0x17, 0x2A)   # slate-950
SLATE700 = RGBColor(0x33, 0x41, 0x55)   # slate-700
SLATE400 = RGBColor(0x94, 0xA3, 0xB8)   # slate-400
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
GREEN    = RGBColor(0x10, 0xB9, 0x81)   # emerald-500


def set_cell_background(cell, hex_color: str):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4F46E5')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_page_break(doc):
    doc.add_page_break()


def styled_heading(doc, text, level=1, color=None):
    """Add a heading with custom colour."""
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    h.paragraph_format.space_after  = Pt(6)
    for run in h.runs:
        run.font.color.rgb = color or INDIGO
        if level == 1:
            run.font.size = Pt(20)
        elif level == 2:
            run.font.size = Pt(15)
        else:
            run.font.size = Pt(13)
    return h


def body_para(doc, text, bold=False, color=None, size=11, space_after=6):
    """Add a normal body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def bullet(doc, text, level=0):
    """Add a bullet point paragraph."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def add_screenshot(doc, filename, caption, width=6.0):
    """Embed a screenshot with a caption beneath it."""
    img_path = IMAGES / filename
    if not img_path.exists():
        body_para(doc, f"[Screenshot not available: {filename}]", color=SLATE400)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    r = cap.add_run(f"Figure: {caption}")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = SLATE400


def add_kv_table(doc, rows: list[tuple[str, str]]):
    """Add a two-column key-value info table."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.5)
    for i, (k, v) in enumerate(rows):
        kc = table.cell(i, 0)
        vc = table.cell(i, 1)
        set_cell_background(kc, '4F46E5')
        kp = kc.paragraphs[0]
        kr = kp.add_run(k)
        kr.bold = True
        kr.font.color.rgb = WHITE
        kr.font.size = Pt(10)

        vp = vc.paragraphs[0]
        vr = vp.add_run(v)
        vr.font.size = Pt(10)
    doc.add_paragraph()


def build_title_page(doc):
    """Build the professional title page."""
    # Big logo-style spacer
    for _ in range(3):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = title_p.add_run("Engage AI")
    r1.bold = True
    r1.font.size = Pt(42)
    r1.font.color.rgb = INDIGO

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitle_p.add_run("AI YouTube Creator Engagement Suite & Live Assistant Manager")
    r2.font.size = Pt(16)
    r2.font.color.rgb = VIOLET
    r2.italic = True

    doc.add_paragraph()
    add_horizontal_rule(doc)
    doc.add_paragraph()

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_r = meta_p.add_run(
        f"Project Description  ·  Version 1.0  ·  {date.today().strftime('%B %d, %Y')}"
    )
    meta_r.font.size = Pt(11)
    meta_r.font.color.rgb = SLATE400

    for _ in range(6):
        doc.add_paragraph()

    add_page_break(doc)


def build_toc(doc):
    styled_heading(doc, "Table of Contents", level=1)
    toc_items = [
        ("1", "Executive Summary"),
        ("2", "Key Features Overview"),
        ("3", "Application Screens & Screenshots"),
        ("3.1", "Login Page"),
        ("3.2", "Dashboard"),
        ("3.3", "AI Video Analyzer"),
        ("3.4", "YouTube Live Chat Manager"),
        ("3.5", "Comment Categorization Engine"),
        ("3.6", "Comment Decision Engine"),
        ("3.7", "Interactive Analytics"),
        ("3.8", "AI Learning & Prompt Memory"),
        ("3.9", "Admin Control Panel"),
        ("3.10", "Personality Voice Presets"),
        ("3.11", "Settings"),
        ("3.12", "Billing"),
        ("4", "Technical Architecture"),
        ("5", "Database & Data Model"),
        ("6", "API & Integration Specifications"),
        ("7", "Security & Authentication"),
        ("8", "Deployment Guide"),
        ("9", "Developer Onboarding"),
    ]
    for num, label in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"  {num}   {label}")
        r.font.size = Pt(10.5)
        if '.' not in num:
            r.bold = True
            r.font.color.rgb = INDIGO
    add_page_break(doc)


def build_executive_summary(doc):
    styled_heading(doc, "1. Executive Summary", level=1)
    body_para(doc,
        "Engage AI is a production-quality, modular, and fully responsive SaaS web application "
        "engineered to help YouTube creators, brands, and agencies manage and automate their "
        "audience interaction at scale. The platform delivers a complete YouTube community "
        "management solution by unifying video comment processing, context-aware AI reply "
        "generation, a real-time live chat overlay, analytics, and a continuous prompt "
        "calibration feedback loop — all within a single, beautifully designed interface."
    )
    body_para(doc,
        "Built on a modern React 19 + Vite + TypeScript foundation with Tailwind CSS for "
        "adaptive dark/light theming, Engage AI is production-ready and designed for seamless "
        "scalability. Authentication and user management are handled by Supabase, while "
        "all AI interactions are driven by a calibrated prompt system supporting nine distinct "
        "personality voice presets."
    )
    body_para(doc,
        "The platform was developed across 11 distinct engineering phases, each incrementally "
        "adding sophisticated capabilities — from video metadata crawling and sentiment analysis "
        "through to admin control panels, quota monitoring, and broadcast notifications."
    )

    add_kv_table(doc, [
        ("Project Name",    "Engage AI"),
        ("Type",            "SaaS Web Application (Single Page App)"),
        ("Version",         "0.0.0 (MVP)"),
        ("Frontend Stack",  "React 19, TypeScript, Vite 8, Tailwind CSS 4"),
        ("Auth Provider",   "Supabase (with secure Mock Auth fallback)"),
        ("Build Tool",      "Vite with lazy routing and code splitting"),
        ("Code Linting",    "Oxlint (optimised code-quality auditing)"),
        ("Target Users",    "YouTube Creators, Brands, Agencies"),
    ])
    add_page_break(doc)


def build_features(doc):
    styled_heading(doc, "2. Key Features Overview", level=1)

    features = [
        ("AI Video Analyzer (Phases 3 & 5)",
         "Crawls video metadata including transcripts, view counts, and upload dates. "
         "Segregates viewer comments into four smart buckets: Questions, Praise, Feedback, and Spam — "
         "giving creators an immediate, structured overview of audience sentiment.",
         ["Automated metadata extraction from YouTube video URLs",
          "Smart comment bucketing: Questions / Praise / Feedback / Spam",
          "Transcript-aware analysis for deeper context understanding",
          "Visual comment stream with per-category filtering controls"]),

        ("Comment Intelligence Engine (Phase 6)",
         "A deep analysis layer that processes every fetched comment for sentiment polarity "
         "(Positive, Neutral, Negative) and assigns priority tags to surface the most important "
         "interactions first. Includes automated spam bot detection, malicious link flagging, "
         "and toxic comment suppression to prevent auto-replies from being sent to bad actors.",
         ["Sentiment analysis: Positive / Neutral / Negative tags",
          "Priority tagging to surface high-value interactions first",
          "Spam bot and malicious link detection",
          "Toxic comment suppression to prevent unintended auto-replies"]),

        ("Calibrated Reply Presets (Phase 7)",
         "Nine distinct personality voice presets allow creators to tailor their AI-generated "
         "replies to perfectly match their channel's brand identity. Strict custom moderation "
         "filters are enforced at the output layer, including link stripping, email redacting, "
         "and release date placeholder replacements.",
         ["9 personality voices: Professional, Friendly, Funny, Spiritual, Gaming, Finance, Tech, Education, Minimal",
          "Per-preset tone rules with configurable verbosity",
          "Link stripping and email redaction filters",
          "Release date placeholder substitution in replies"]),

        ("YouTube Live Chat Manager (Phase 8)",
         "A real-time simulated stream comments feed with distinct visual tags for channel "
         "Members, Moderators, and highlighted Super Chats. Hover-based moderation tools provide "
         "instant in-stream actions — including Pin Message, Delete, Timeout, and Ban User — "
         "alongside one-click AI reply draft injection.",
         ["Real-time live comment stream simulation",
          "Visual distinction: Members, Moderators, Super Chats",
          "Hover moderation tools: Pin, Delete, Timeout, Ban",
          "One-click AI-generated reply draft injection"]),

        ("Interactive Analytics (Phase 9)",
         "Comprehensive visual reporting with growth area curves (toggling Daily/Weekly/Monthly), "
         "sentiment distribution rings, horizontal language demographics, and category count charts. "
         "All data is exportable to CSV for spreadsheet analysis or printable PDF for presentations.",
         ["Growth curves with Daily / Weekly / Monthly toggle",
          "Sentiment distribution ring charts",
          "Language demographics horizontal bar charts",
          "Category count charts and comment volume metrics",
          "One-click CSV export and printable PDF report generation"]),

        ("AI Learning & Prompt Memory (Phase 10)",
         "A continuous improvement loop that intercepts every creator edit made to an AI-generated "
         "reply and prompts a micro-rating survey with an adjustment reason. The system automatically "
         "formulates and appends refined prompt instructions to improve the quality of all future "
         "AI completions over time.",
         ["Intercepts creator edits to AI-generated replies",
          "Micro-rating surveys with contextual adjustment reasons",
          "Automatic prompt instruction refinement and appending",
          "Persistent prompt memory improves reply quality over time"]),

        ("Admin Control Panel (Phase 11)",
         "A consolidated operations dashboard for platform administrators, providing full "
         "visibility into API quota consumption, user account management, and system-wide "
         "configuration.",
         ["YouTube API Daily Quota gauge with usage visualisation",
          "User account status manager: upgrade plans, suspend access",
          "Global feature flags configuration panel",
          "System broadcast notification publisher"]),
    ]

    for title, desc, bullets_list in features:
        styled_heading(doc, title, level=2, color=VIOLET)
        body_para(doc, desc)
        for b in bullets_list:
            bullet(doc, b)
        doc.add_paragraph()

    add_page_break(doc)


def build_screenshots(doc):
    styled_heading(doc, "3. Application Screens & Screenshots", level=1)
    body_para(doc,
        "The following section presents high-fidelity screenshots of every screen in the Engage AI "
        "application, captured from the running production build at 1400×900 viewport resolution."
    )
    doc.add_paragraph()

    screens = [
        ("3.1", "Login Page",                     "login.png",
         "The Login Page features a stunning glassmorphism card set against animated gradient "
         "blobs on a slate background. It supports both Sign In and Sign Up modes with a tab "
         "switcher, live error validation, password visibility toggle, and rate-limit protection. "
         "In development environments without Supabase credentials, the app seamlessly operates "
         "in Mock Auth mode, accepting any credentials."),

        ("3.2", "Dashboard",                       "dashboard.png",
         "The Dashboard provides a central command centre with key metric cards (total comments "
         "analysed, replies sent, engagement rate, and pending reviews), a recent activity feed, "
         "quick-action shortcuts to all major features, and a personalised welcome header showing "
         "the creator's profile."),

        ("3.3", "AI Video Analyzer",               "analyze.png",
         "Creators paste a YouTube video URL to trigger an automated analysis pipeline. The page "
         "displays extracted video metadata (thumbnail, title, views, publish date) alongside a "
         "live-updating comment stream sorted into smart categories — Questions, Praise, Feedback, "
         "and Spam — each with sentiment badges and individual reply action buttons."),

        ("3.4", "YouTube Live Chat Manager",        "live_chat.png",
         "A simulated real-time live comment feed displays incoming stream messages with distinct "
         "colour-coded badges for Members, Moderators, and Super Chat supporters. Hovering any "
         "message reveals an inline moderation toolbar (Pin, Delete, Timeout, Ban) and a one-click "
         "AI reply suggestion injection button."),

        ("3.5", "Comment Categorization Engine",    "categorize.png",
         "This dedicated view surfaces all previously crawled comments in a filterable, sortable "
         "table. Creators can toggle between category views (Questions, Praise, Feedback, Spam), "
         "apply sentiment filters, and bulk-approve or dismiss AI reply drafts before publishing."),

        ("3.6", "Comment Decision Engine",          "decision_engine.png",
         "The Decision Engine presents an AI-powered triage workflow that scores each comment's "
         "reply priority. Creators can configure decision thresholds, review the auto-flagged "
         "comments for human review, and approve batches of low-risk auto-replies in one click."),

        ("3.7", "Interactive Analytics",            "analytics.png",
         "A rich analytics dashboard featuring a togglable area growth curve (Daily/Weekly/Monthly), "
         "a sentiment distribution donut chart, a horizontal language demographics bar chart, and "
         "category count rings. Export controls allow one-click CSV data download and printable "
         "PDF report generation."),

        ("3.8", "AI Learning & Prompt Memory",      "ai_learning.png",
         "Every time a creator edits an AI-generated reply, this page records the edit, presents "
         "a 5-star micro-rating survey, and asks for a brief adjustment reason. The collected "
         "feedback is automatically synthesised into new prompt directives that are silently "
         "appended to the AI system prompt, continuously improving reply quality."),

        ("3.9", "Admin Control Panel",              "admin_panel.png",
         "The Admin Panel consolidates platform operations management into a single pane: a "
         "circular YouTube API Daily Quota gauge, a user account table with plan upgrade and "
         "suspend controls, a global feature flags toggle matrix, and a broadcast notification "
         "composer for sending system-wide messages to all users."),

        ("3.10", "Personality Voice Presets",       "personality.png",
         "Nine distinct AI personality cards (Professional, Friendly, Funny, Spiritual, Gaming, "
         "Finance, Tech, Education, Minimal) are available for creators to select. Each preset "
         "has a configurable tone description, verbosity slider, and live reply sample preview. "
         "Custom moderation filter toggles (link stripping, email redaction) are also managed here."),

        ("3.11", "Settings",                        "settings.png",
         "The Settings page allows creators to manage their profile information, connected YouTube "
         "channel, notification preferences, API key configuration, and data privacy controls — "
         "all presented in a clean tabbed interface."),

        ("3.12", "Billing",                         "billing.png",
         "The Billing page shows the creator's current subscription plan with a feature comparison "
         "table, usage metrics against plan limits, and upgrade/downgrade controls with plan "
         "pricing cards."),
    ]

    for num, title, filename, description in screens:
        styled_heading(doc, f"{num}  {title}", level=2, color=INDIGO)
        body_para(doc, description)
        add_screenshot(doc, filename, title, width=6.2)

    add_page_break(doc)


def build_architecture(doc):
    styled_heading(doc, "4. Technical Architecture", level=1)
    body_para(doc,
        "Engage AI is architected as a modular Single Page Application (SPA) following a strict "
        "separation of concerns between presentation, business logic, and data access layers."
    )

    styled_heading(doc, "Frontend Layer", level=2, color=VIOLET)
    bullet(doc, "React 19 with concurrent rendering features for smooth, responsive UI updates")
    bullet(doc, "TypeScript 6 for end-to-end type safety across all components and utility functions")
    bullet(doc, "Vite 8 as the build tool with lightning-fast HMR in development and optimised chunking in production")
    bullet(doc, "Tailwind CSS 4 for a custom dark/light theme with a consistent design token system")
    bullet(doc, "React Router DOM v7 for declarative, code-split lazy routing")

    styled_heading(doc, "State Management", level=2, color=VIOLET)
    bullet(doc, "React Context API for global Auth and Dashboard state slices")
    bullet(doc, "Local component state (useState/useReducer) for all page-specific interactions")
    bullet(doc, "Custom hooks (useAuth, useDashboard) encapsulate reusable stateful logic")

    styled_heading(doc, "Data Layer", level=2, color=VIOLET)
    bullet(doc, "Supabase (PostgreSQL) for production authentication and user session management")
    bullet(doc, "Mock Auth mode for local development without requiring Supabase credentials")
    bullet(doc, "Zod for runtime schema validation of all external and localStorage data")

    styled_heading(doc, "Code Quality", level=2, color=VIOLET)
    bullet(doc, "Oxlint for fast, Rust-based code quality linting across the entire codebase")
    bullet(doc, "Strict TypeScript configuration with no implicit any and full module resolution")
    bullet(doc, "Pre-build environment secret scanning script to prevent credential leakage")

    add_page_break(doc)


def build_database(doc):
    styled_heading(doc, "5. Database & Data Model", level=1)
    body_para(doc,
        "Engage AI uses Supabase (PostgreSQL) as its primary backend with Row Level Security (RLS) "
        "enforcing strict per-user data isolation. In development, a localStorage-based mock "
        "client simulates the same data shape for seamless offline development."
    )

    styled_heading(doc, "Core Entities", level=2, color=VIOLET)
    entities = [
        ("users",           "Extends Supabase auth.users with creator profile metadata, connected channel info, and plan tier"),
        ("videos",          "Crawled YouTube video metadata: URL, title, thumbnail, view count, publish date, transcript hash"),
        ("comments",        "Raw YouTube comments with category tag, sentiment score, priority score, spam flag, and reply status"),
        ("reply_drafts",    "AI-generated reply drafts with personality preset, generation timestamp, and approval status"),
        ("prompt_memory",   "Creator edit intercepts: original reply, edited reply, rating, adjustment reason, generated directive"),
        ("live_sessions",   "YouTube live stream sessions with associated chat messages and moderation action logs"),
        ("analytics_cache", "Pre-aggregated daily/weekly/monthly metric snapshots for fast analytics rendering"),
    ]
    for name, desc in entities:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"  {name}:  ")
        r1.bold = True
        r1.font.color.rgb = VIOLET
        r1.font.size = Pt(10.5)
        r2 = p.add_run(desc)
        r2.font.size = Pt(10.5)

    add_page_break(doc)


def build_api(doc):
    styled_heading(doc, "6. API & Integration Specifications", level=1)
    body_para(doc,
        "Engage AI integrates with several external services. All API keys are stored "
        "exclusively in server-side environment variables and are never exposed to the client bundle."
    )

    integrations = [
        ("YouTube Data API v3",
         ["Fetch video metadata: snippet, statistics, contentDetails",
          "Retrieve comment threads with author details and publish timestamps",
          "Daily quota: 10,000 units (monitored in Admin Panel)",
          "Scopes: youtube.readonly, youtubepartner"]),
        ("Supabase REST & Realtime API",
         ["Auth: signUp, signInWithPassword, onAuthStateChange, signOut",
          "Database: CRUD via auto-generated REST endpoints with RLS",
          "Realtime subscriptions for live comment stream simulation",
          "Storage: avatar and thumbnail asset hosting"]),
        ("AI Completion API (Prompt System)",
         ["System prompt construction with active personality preset rules",
          "Moderation filter post-processing (link stripping, email redacting)",
          "Prompt memory directive injection from AI Learning module",
          "Context window: video metadata + comment text + creator instructions"]),
    ]

    for service, points in integrations:
        styled_heading(doc, service, level=2, color=VIOLET)
        for pt in points:
            bullet(doc, pt)
        doc.add_paragraph()

    add_page_break(doc)


def build_security(doc):
    styled_heading(doc, "7. Security & Authentication", level=1)
    body_para(doc,
        "Security is a first-class concern throughout the Engage AI architecture. "
        "Multiple complementary layers protect user data and prevent abuse."
    )

    measures = [
        ("HTTPS Enforcement",          "The app enforces HTTPS at runtime via enforceHttps(), immediately redirecting any HTTP request."),
        ("CSRF Token Protection",       "A cryptographically random CSRF token is generated on login (bootstrapSession()) and stored in sessionStorage. All state-mutating operations validate this token before execution."),
        ("Session Expiry Watcher",      "In Mock Auth mode, a 60-second polling interval checks for session expiry (isSessionExpired()) and automatically tears down the session if it has elapsed."),
        ("Login Rate Limiting",         "checkAndRecordLoginAttempt() enforces a sliding-window rate limit on login submissions, preventing brute-force attacks."),
        ("Input Validation (Zod)",      "All data read from localStorage, external APIs, or URL parameters is validated against strict Zod schemas before use in the application."),
        ("Environment Secret Scanning", "A pre-build Node.js script (check-env-secrets.js) scans all source files for accidentally committed secrets, blocking the production build if any are detected."),
        ("Supabase RLS",                "Row Level Security policies on all Supabase tables ensure users can only read and write their own data, even if an API key were somehow compromised."),
        ("Password Policy",             "Sign-up enforces a minimum password length of 8 characters. Supabase handles bcrypt hashing server-side."),
    ]

    for measure, desc in measures:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        r1 = p.add_run(f"  {measure}: ")
        r1.bold = True
        r1.font.color.rgb = INDIGO
        r1.font.size = Pt(10.5)
        r2 = p.add_run(desc)
        r2.font.size = Pt(10.5)

    add_page_break(doc)


def build_deployment(doc):
    styled_heading(doc, "8. Deployment Guide", level=1)

    styled_heading(doc, "Environment Variables", level=2, color=VIOLET)
    body_para(doc, "Create a .env.local file in the project root with the following variables:")

    env_vars = [
        ("VITE_SUPABASE_URL",      "Your Supabase project URL (https://xxxxx.supabase.co)"),
        ("VITE_SUPABASE_ANON_KEY", "Your Supabase public anon key"),
        ("VITE_YOUTUBE_API_KEY",   "Your Google YouTube Data API v3 key"),
    ]
    for k, v in env_vars:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"  {k}")
        r1.bold = True
        r1.font.name = 'Courier New'
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = VIOLET
        r2 = p.add_run(f"  —  {v}")
        r2.font.size = Pt(9.5)

    styled_heading(doc, "Local Development", level=2, color=VIOLET)
    cmds = [
        ("npm install",   "Install all npm dependencies"),
        ("npm run dev",   "Boot the Vite HMR development server at http://localhost:5173"),
        ("npm run lint",  "Run Oxlint for code quality checks"),
        ("npm run build", "Compile TypeScript and produce optimised production bundle in /dist"),
    ]
    for cmd, desc in cmds:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"  $ {cmd}")
        r1.font.name = 'Courier New'
        r1.font.size = Pt(9.5)
        r1.bold = True
        r1.font.color.rgb = INDIGO
        r2 = p.add_run(f"  —  {desc}")
        r2.font.size = Pt(9.5)

    styled_heading(doc, "Static Hosting (Firebase)", level=2, color=VIOLET)
    bullet(doc, "A firebase.json configuration is included for direct deployment to Firebase Hosting")
    bullet(doc, "Run npm run build, then firebase deploy --only hosting")
    bullet(doc, "All routes are rewritten to index.html for SPA compatibility")
    bullet(doc, "Vite output directory /dist is set as the public root")

    add_page_break(doc)


def build_developer_guide(doc):
    styled_heading(doc, "9. Developer Onboarding", level=1)

    styled_heading(doc, "Project Structure", level=2, color=VIOLET)
    structure = [
        ("src/pages/",      "One file per top-level route. Lazy-loaded via React.lazy()"),
        ("src/components/", "Reusable UI components organised by feature domain"),
        ("src/contexts/",   "Global React Context providers: AuthContext, DashboardContext"),
        ("src/hooks/",      "Custom React hooks encapsulating reusable stateful logic"),
        ("src/services/",   "API service modules (YouTube, Supabase, AI completion)"),
        ("src/utils/",      "Pure utility functions: validation, auth security, clipboard, URL helpers"),
        ("src/supabase/",   "Supabase client initialisation and configuration"),
        ("src/layouts/",    "DashboardLayout wrapper with sidebar navigation and top bar"),
        ("scripts/",        "Build-time Node.js utility scripts"),
        ("public/",         "Static assets served as-is by Vite"),
    ]
    for path, desc in structure:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"  {path}")
        r1.font.name = 'Courier New'
        r1.font.size = Pt(9.5)
        r1.bold = True
        r1.font.color.rgb = VIOLET
        r2 = p.add_run(f"  —  {desc}")
        r2.font.size = Pt(9.5)

    styled_heading(doc, "Adding a New Page", level=2, color=VIOLET)
    steps = [
        "Create the page component file in src/pages/MyNewPage.tsx",
        "Export the component as a named export: export const MyNewPage: React.FC = () => { ... }",
        "Add a lazy import in src/App.tsx: const MyNewPage = lazy(() => import('./pages/MyNewPage'))",
        "Register the route inside the protected Routes block in App.tsx",
        "Add the navigation entry to src/layouts/DashboardLayout.tsx sidebar config",
    ]
    for i, step in enumerate(steps, 1):
        bullet(doc, f"Step {i}: {step}")

    styled_heading(doc, "TypeScript Conventions", level=2, color=VIOLET)
    bullet(doc, "All props interfaces are defined above the component in the same file")
    bullet(doc, "Prefer named exports over default exports for all components")
    bullet(doc, "Use Zod schemas for any data that comes from external sources or localStorage")
    bullet(doc, "Context values must be fully typed — no implicit any in context definitions")


def main():
    print("📄  Building Engage AI Project Description document...")

    doc = Document()

    # ─── Page margins ────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # ─── Default font ────────────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ─── Build all sections ──────────────────────────────────────────────────
    build_title_page(doc)
    build_toc(doc)
    build_executive_summary(doc)
    build_features(doc)
    build_screenshots(doc)
    build_architecture(doc)
    build_database(doc)
    build_api(doc)
    build_security(doc)
    build_deployment(doc)
    build_developer_guide(doc)

    # ─── Save ────────────────────────────────────────────────────────────────
    doc.save(str(OUTPUT))
    size_kb = OUTPUT.stat().st_size // 1024
    print(f"✅  Saved: {OUTPUT}  ({size_kb} KB)")


if __name__ == '__main__':
    main()
